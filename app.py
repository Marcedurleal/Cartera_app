import streamlit as st
import pandas as pd
import io
import logging
from datetime import datetime
from docx import Document

# --------------------- CONFIGURACIÓN INICIAL ---------------------
st.set_page_config(page_title="Cartera por Torre", layout="wide")

st.markdown("""
<h1 style='text-align: center;'>📄 Generador de Cartera por Torre</h1>
<p style='text-align: center; font-size: 18px;'>Sube el archivo Excel de Cartera y genera el reporte consolidado en Word por Torre.</p>
""", unsafe_allow_html=True)

# Configurar logs
log_messages = []

def log(msg, level="info"):
    """Registra logs visibles en la barra lateral."""
    log_messages.append((level, msg))
    if level == "error":
        logging.error(msg)
    elif level == "warning":
        logging.warning(msg)
    else:
        logging.info(msg)

def format_currency(val):
    """Formatea valores numéricos como moneda ($126.850 o -$126.850)."""
    try:
        val = float(val)
        if val < 0:
            return f"-${abs(val):,.0f}".replace(",", ".")
        else:
            return f"${val:,.0f}".replace(",", ".")
    except (ValueError, TypeError):
        return "$0"


# -------------------------- CARGA DEL ARCHIVO ---------------------
st.subheader("📥 Cargar archivo")
uploaded_file = st.file_uploader("Selecciona un archivo Excel (.xlsx)", type=["xlsx"])

if uploaded_file is not None:

    st.toast("Procesando archivo...", icon="⏳")

    # 1. Leer únicamente la hoja CARTERA
    try:
        df_cartera = pd.read_excel(uploaded_file, sheet_name='CARTERA')
        log("Hoja 'CARTERA' cargada correctamente ✔")
    except Exception as e:
        st.error("❌ Error al leer la hoja 'CARTERA'. Verifica que exista en el archivo Excel.")
        log(f"Error cargando la hoja CARTERA: {str(e)}", "error")
        st.stop()

    # Validar columnas obligatorias
    required_cols = ["codigo", "interior", "total"]
    missing_cols = [c for c in required_cols if c not in df_cartera.columns]

    if missing_cols:
        st.error(f"❌ La hoja CARTERA no contiene las columnas necesarias: {missing_cols}")
        log(f"Columnas faltantes: {missing_cols}", "error")
        st.stop()

    log("Validación de columnas completada ✔")

    # --------------------- CONFIGURACIÓN Y ENTRADA DE FECHA ---------------------
    st.subheader("📅 Parámetros del Informe")
    Fecha_corte_str = st.text_input("Fecha de corte (dd/mm/aaaa):", placeholder="Ej: 31/08/2026")

    # --------------------- GENERACIÓN DEL REPORTE WORD ---------------------
    if st.button("📄 Generar Reporte Word"):

        if not Fecha_corte_str:
            st.error("⚠ Debes ingresar la fecha de corte.")
            st.stop()

        try:
            Fecha_corte = datetime.strptime(Fecha_corte_str, "%d/%m/%Y").date()
        except Exception:
            st.error("❌ Formato de fecha inválido. Debe ser dd/mm/aaaa (Ej: 15/08/2026).")
            st.stop()

        try:
            # ---------------- TRATAMIENTO DE DATOS ----------------
            df_proc = df_cartera.copy()

            # Extraer y formatear valor de c_13050506 (Fachadas)
            if "c_13050506" in df_proc.columns:
                fachadas_val = df_proc["c_13050506"].fillna(0)
            else:
                fachadas_val = 0

            # 1. Cálculo de Administración y parqueaderos = total - c_13050506
            # (Se permite que el resultado sea negativo en caso de haber anticipos/saldos a favor)
            df_proc["Administración y parqueaderos"] = df_proc["total"].fillna(0) - fachadas_val

            # 2. Asignar/Renombrar columna 'Fachadas'
            if "c_13050506" in df_proc.columns:
                df_proc.rename(columns={"c_13050506": "Fachadas"}, inplace=True)
            else:
                df_proc["Fachadas"] = 0

            df_proc["Fachadas"] = df_proc["Fachadas"].fillna(0)

            # 3. FILTRADO: total > 1000 Y 'interior' no sea 0, "0", ni nulo/vacío
            cond_total = df_proc["total"] > 1000
            cond_interior_valido = (
                df_proc["interior"].notna() &
                (~df_proc["interior"].isin([0, "0", 0.0, "0.0"])) &
                (df_proc["interior"].astype(str).str.strip() != "")
            )

            df_filtered = df_proc[cond_total & cond_interior_valido].copy()

            if df_filtered.empty:
                st.warning("No se encontraron registros de morosos válidos (total > $1.000 e interior distinto de 0 o nulo).")
                st.stop()

            # ---------------- CONSTRUCCIÓN DEL DOCUMENTO WORD ----------------
            document = Document()
            unique_towers = df_filtered["interior"].unique()

            first_tower = True

            for tower in unique_towers:

                # Limpieza del nombre de la Torre (convierte '12.0' -> '12')
                try:
                    tower_clean = str(int(float(tower)))
                except (ValueError, TypeError):
                    tower_clean = str(tower).strip()

                if not first_tower:
                    document.add_page_break()
                else:
                    first_tower = False

                df_tower = df_filtered[df_filtered["interior"] == tower]

                # Encabezado por Torre
                document.add_heading(f"MOROSOS TORRE {tower_clean}", level=1)
                document.add_paragraph(
                    f"A continuación se relacionan los morosos de la torre {tower_clean} "
                    f"con corte a {Fecha_corte.strftime('%d/%m/%Y')}"
                )

                # Tabla con 4 columnas
                table = document.add_table(rows=1, cols=4)
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Código"
                hdr_cells[1].text = "Administración y parqueaderos"
                hdr_cells[2].text = "Fachadas"
                hdr_cells[3].text = "Total"

                for _, row in df_tower.iterrows():
                    row_cells = table.add_row().cells

                    # Formatear el código si es numérico (quitar .0)
                    try:
                        codigo_str = str(int(float(row["codigo"])))
                    except (ValueError, TypeError):
                        codigo_str = str(row["codigo"])

                    row_cells[0].text = codigo_str
                    row_cells[1].text = format_currency(row["Administración y parqueaderos"])
                    row_cells[2].text = format_currency(row["Fachadas"])
                    row_cells[3].text = format_currency(row["total"])

            # Guardar en buffer de memoria
            buffer_word = io.BytesIO()
            document.save(buffer_word)
            buffer_word.seek(0)

            st.success("✔ Documento Word generado con éxito.")

            st.download_button(
                label="📥 Descargar Word – Cartera por Torre",
                data=buffer_word,
                file_name=f"Cartera_por_Torre_{Fecha_corte.strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error("❌ Ocurrió un error inesperado al procesar el archivo o generar el Word.")
            st.write(f"Detalle del error: {str(e)}")

# ---------------- LOGS EN LA BARRA LATERAL ----------------
with st.sidebar:
    st.header("📝 Logs del proceso")
    for level, msg in log_messages:
        if level == "error":
            st.error("❌ " + msg)
        elif level == "warning":
            st.warning("⚠ " + msg)
        else:
            st.info("ℹ " + msg)
